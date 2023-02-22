from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx
from docx.shared import Pt

import mtk2


def get_or_add_spacing(rPr):
    spacings = rPr.xpath("./w:spacing")
    if spacings:
        return spacings[0]
    spacing = OxmlElement("w:spacing")
    rPr.insert_element_before(
        spacing,
        *(
            "w:w",
            "w:kern",
            "w:position",
            "w:sz",
            "w:szCs",
            "w:highlight",
            "w:u",
            "w:effect",
            "w:bdr",
            "w:shd",
            "w:fitText",
            "w:vertAlign",
            "w:rtl",
            "w:cs",
            "w:em",
            "w:lang",
            "w:eastAsianLayout",
            "w:specVanish",
            "w:oMath",
        ),
    )
    return spacing


def run_set_spacing(run, value: int):
    rPr = run._r.get_or_add_rPr()
    spacing = get_or_add_spacing(rPr)
    spacing.set(qn('w:val'), str(value))


def add_doc():
    document = docx.Document("variant10.docx")

    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    for i in document.paragraphs:
        doc.add_paragraph()

    for p in range(len(document.paragraphs)):
        for t in document.paragraphs[p].text:
            doc.paragraphs[p].add_run(t)

    return doc


def hide_text():
    doc = add_doc()

    text = "Баклажан на стебле дыни не вырастет"
    code_text = mtk2.MTK2_code(text)

    index = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if code_text[index] == '1':
                run_set_spacing(run, 2)
            if index < len(code_text)-1:
                index += 1

    path = "secret_message.docx"
    doc.save(path)

    print()
    print(text, "---> MTK2 encode:")
    print(code_text)
    print("\nСодержимое файла:")
    print("-" * 100)

    for paragraph in doc.paragraphs:
        print(paragraph.text)

    print("-" * 100)
    print("\nФайл", path, "сохранен")


hide_text()
